import os
from telegram import Update
from telegram.ext import ApplicationBuilder, CommandHandler, MessageHandler, ContextTypes, filters
from docx import Document

def load_tasks():
    doc = Document("Задания.docx")
    tasks = []
    current = ""
    for p in doc.paragraphs:
        if p.text.startswith("📘 Задание"):
            if current:
                tasks.append(current.strip())
            current = p.text.strip()
        else:
            current += "\n" + p.text.strip()
    if current:
        tasks.append(current.strip())
    return tasks

def load_answers():
    doc = Document("Ответы.docx")
    answers = []
    current = ""
    for p in doc.paragraphs:
        if p.text.startswith("📘 Задание"):
            if current:
                answers.append(current.strip())
            current = ""
        else:
            current += "\n" + p.text.strip()
    if current:
        answers.append(current.strip())
    return answers

tasks = load_tasks()
answers = load_answers()
user_states = {}

async def start(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    user_states[user_id] = 0
    await update.message.reply_text(f"Привет! Вот задание №1:\n\n{tasks[0]}")

async def handle_answer(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user_id = update.effective_user.id
    q = user_states.get(user_id, 0)
    user_text = update.message.text.strip()
    correct = answers[q]

    if user_text.lower() in correct.lower():
        reply = "✅ Верно!"
    else:
        reply = f"❌ Неверно. Правильный ответ:\n{correct}"

    q += 1
    if q < len(tasks):
        user_states[user_id] = q
        reply += f"\n\nСледующее задание:\n\n{tasks[q]}"
    else:
        reply += "\n\n🎉 Это было последнее задание!"
        user_states.pop(user_id, None)

    await update.message.reply_text(reply)

async def main():
    app = ApplicationBuilder().token(os.environ["TOKEN"]).build()
    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, handle_answer))

    await app.initialize()
    await app.start()
    print("🤖 Бот запущен.")
    await app.updater.start_polling()
    await app.updater.idle()
    await app.stop()
    await app.shutdown()

if __name__ == "__main__":
    import asyncio
    loop = asyncio.get_event_loop()
    loop.create_task(main())
    loop.run_forever()
